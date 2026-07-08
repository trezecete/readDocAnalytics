from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.analysis.models import AnalysisReport, Finding
from app.reports.presenter import (
    FINDING_TYPE_LABELS,
    SEVERITY_LABELS,
    SEVERITY_ORDER,
    build_report_presenter,
)


def build_docx_report(report: AnalysisReport) -> bytes:
    presenter = build_report_presenter(report)
    document = Document()
    _configure_document(document)

    title = document.add_heading("Relatório executivo de análise", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(report.document_title)
    document.add_paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    document.add_heading("Resumo da proposta", level=1)
    document.add_paragraph(report.executive_summary)

    document.add_heading("Principais pontos de atenção", level=1)
    summary_table = document.add_table(rows=1, cols=4)
    summary_table.style = "Table Grid"
    headers = summary_table.rows[0].cells
    headers[0].text = "Críticos"
    headers[1].text = "Altos"
    headers[2].text = "Lacunas"
    headers[3].text = "Modo"
    row = summary_table.add_row().cells
    row[0].text = str(presenter.severity_counts["critica"])
    row[1].text = str(presenter.severity_counts["alta"])
    row[2].text = str(len(presenter.gap_findings))
    row[3].text = presenter.mode_label

    document.add_paragraph(f"Próximo passo recomendado: {presenter.primary_next_step}")

    if report.coverage:
        document.add_heading("Como a análise foi conduzida", level=1)
        document.add_paragraph(
            f"Foram considerados {report.coverage.document_chars} caracteres, com "
            f"{report.coverage.text_scanned_percent:.0f}% do texto varrido ou indexado."
        )
        document.add_paragraph(
            "Categorias verificadas: " + ", ".join(report.coverage.categories_checked)
        )
        document.add_paragraph(report.coverage.caveat)

    for severity in SEVERITY_ORDER:
        findings = presenter.grouped_findings.get(severity, [])
        if not findings:
            continue
        document.add_heading(SEVERITY_LABELS[severity], level=1)
        for finding in findings:
            _add_finding(document, finding)

    document.add_heading("Plano de ação recomendado", level=1)
    for recommendation in report.recommendations:
        document.add_paragraph(recommendation, style="List Bullet")

    document.add_heading("Limitações da análise", level=1)
    for limitation in report.limitations:
        document.add_paragraph(limitation, style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(35, 31, 32)

    for style_name in ["Heading 1", "Heading 2", "Title"]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = RGBColor(102, 0, 153)


def _add_finding(document: Document, finding: Finding) -> None:
    document.add_heading(finding.title, level=2)
    paragraph = document.add_paragraph()
    paragraph.add_run("Categoria: ").bold = True
    paragraph.add_run(finding.category)
    paragraph.add_run(" | Tipo: ").bold = True
    paragraph.add_run(FINDING_TYPE_LABELS[finding.finding_type])
    paragraph.add_run(" | Confiança: ").bold = True
    paragraph.add_run(finding.confidence)

    if finding.evidence:
        evidence = document.add_paragraph()
        evidence.add_run("Evidência: ").bold = True
        evidence.add_run(finding.evidence)

    document.add_paragraph(f"Por que importa: {finding.explanation}")
    document.add_paragraph(f"Ação sugerida: {finding.recommendation}")

